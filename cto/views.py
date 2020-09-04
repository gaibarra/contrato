
import io
from pathlib import Path
import os




# Para utilizar algunas de las funciones de la librería
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_BREAK_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT


import docx
import locale


from django.shortcuts import render,redirect, get_list_or_404
from django.views import generic 
from django.views.generic import TemplateView, ListView, CreateView
from django.contrib.messages.views import SuccessMessageMixin

from django.urls import reverse_lazy
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.models import User               
from django.http import HttpResponse

from datetime import datetime, date
from django.contrib import messages
from django.db.models import Q
from django.contrib.auth import authenticate
from bases.views import SinPrivilegios
from bases.views import *
from .models import Departamento, Partes, Contratos, Doctos, Tipocontrato, Requisitos, Valida, Secuencia, Regimen
from .forms import DepartamentoForm, PartesForm, ContratosForm
from bases.views import SinPrivilegios
from django.core.files.storage import FileSystemStorage




class VistaBaseCreate(SuccessMessageMixin,SinPrivilegios,generic.CreateView):
    context_object_name = 'obj'
    success_message="Registro Agregado Satisfactoriamente"

    def form_valid(self, form):
        form.instance.uc = self.request.user
        return super().form_valid(form)




class VistaBaseEdit(SuccessMessageMixin,SinPrivilegios, generic.UpdateView):
    context_object_name = 'obj'
    success_message="Registro Actualizado Satisfactoriamente"

    def form_valid(self, form):
        #form.instance.um = self.request.user.id
        return super().form_valid(form)




class DepartamentoView(SinPrivilegios,generic.ListView):
        model = Departamento
        template_name = "cto:departamento_list.html"
        context_object_name = "obj"
        permission_required = "cto.view_departamento"


class DepartamentoNew(VistaBaseCreate):
    model=Departamento
    template_name="cto/departamento_form.html"
    form_class=DepartamentoForm
    success_url= reverse_lazy("cto:departamento_list")
    permission_required="cto.add_departamento" 

class DepartamentoEdit(VistaBaseEdit):
    model = Departamento
    template_name="cto/departamento_form.html"
    form_class = DepartamentoForm
    success_url= reverse_lazy("cto:departamento_list")
    permission_required="cto.change_departamento"


@login_required(login_url="/login/")
@permission_required("cto.change_departamento",login_url="/login/")
def departamentoInactivar(request,id):
    departamento = Departamento.objects.filter(pk=id).first()

    if request.method=="POST":
        if departamento:
            departamento.estado = not departamento.estado
            departamento.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")
    
    return HttpResponse("FAIL") 

class PartesView(SinPrivilegios,generic.ListView):
    model = Partes
    template_name = "cto:partes_list.html"
    context_object_name = "obj"
    success_url= reverse_lazy("cto:partes_list")
    permission_required = "cto.view_partes"

    def get_queryset(self):
        current_userx = self.request.user.id
        #print(current_userx)
        queryset = Partes.objects.filter(user_id=current_userx)
        for part in queryset:
            xdepa=part.claveDepartamento_id
        
        #print(xdepa)
        querydep = Departamento.objects.all()    
        for depa in querydep:
            
            if depa.claveDepartamento==xdepa:
               xr1=depa.rango1    
               xr2=depa.rango2
             
        
        #print(xr1)
        #print(xr2)
        return Partes.objects.filter(
            Q(estado=True), Q(claveDepartamento__gte=xr1),  Q(claveDepartamento__lte=xr2) 
        )
            
       
                   
    def get_context_data(self, **kwargs):
        # Call the base implementation first to get a context
        context = super(PartesView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
       
        return context      
    
    
    




class PartesNew(VistaBaseCreate):
    model= Partes
    template_name="cto/partes_form.html"
    form_class=PartesForm
    success_url= reverse_lazy("cto:partes_list")
    permission_required="cto.add_partes" 

class PartesEdit(VistaBaseEdit):
    model = Partes
    template_name="cto/partes_form.html"
    form_class = PartesForm
    success_url= reverse_lazy("cto:partes_list")
    permission_required="cto.change_partes"


@login_required(login_url="/login/")
@permission_required("cto.change_partes",login_url="/login/")
def partesInactivar(request,id):
    partes = Partes.objects.filter(pk=id).first()

    if request.method=="POST":
        if partes:
            partes.estado = not partes.estado
            partes.save()
            return HttpResponse("OK")
        return HttpResponse("FAIL")
    
    return HttpResponse("FAIL")   
class ContratosView(SinPrivilegios, generic.ListView):
    model = Contratos
    template_name = "cto/contrato_list.html"
    context_object_name = "obj"
    success_url= reverse_lazy("cto:contrato_list")
    permission_required="cto.view_contratos"
        
    def get_queryset(self):
        
        current_userx = self.request.user.id
        #conditions = dict(current_user=current_userx, uc_id=self.request.user) 
        #queryset = queryset.filter(**conditions)
        return Contratos.objects.all()
        #return Contratos.objects.filter(
        #    Q(current_user=current_userx) | Q(uc_id=self.request.user)
        #)
        #return SpyorEnc.objects.filter(
        #    Q(current_user=current_userx) | Q(uc_id=self.request.user) | Q(el_jefe=current_userx)
        #)
        
    
    def get_context_data(self, **kwargs):
        # Call the base implementation first to get a context
        context = super(ContratosView, self).get_context_data(**kwargs)
        # Get the blog from id and add it to the context
        context['some_data'] = Partes.objects.all()
        context['some_data2'] = Departamento.objects.all()
        return context                   

@login_required(login_url='/login/')
@permission_required('cto.add_contratos', login_url='bases:sin_privilegios')
def contratos2(request,contrato_id=None):
    
    template_name='cto/contrato.html'
    detalle = {}
    secuencia_data = {}
    xUsuario = (request.user.id)
    print(xUsuario)
   
    partes = Partes.objects.filter(estado=True, user=xUsuario)    
    
    p=partes.first()
    print (p)
    
    
    d1 = p.claveDepartamento
    dx = p.claveDepartamento_id
    rfcparte = (p.rfc)
    print(rfcparte)
    
    print(d1)
   
   
    
    requisitos = Requisitos.objects.filter(estado=True)
    departamentos =Departamento.objects.filter(estado=True, claveDepartamento=dx)
    departamentos2 =Departamento.objects.filter(estado=True)
    d2=departamentos.first()
    d3= (d2.claveDepartamento)
    r1= (d2.rango1)
    r2= (d2.rango2)
    secuencia = (d2.f001)

    print (d3)
    print(secuencia)
    
    
    partes3 =Partes.objects.filter(Q(estado=True),  Q(claveDepartamento_id__gte=r1), Q(claveDepartamento_id__lte=r2) ).order_by('nombreParte')
    
    
    partes2 =Partes.objects.filter(estado=True).order_by('nombreParte')
    print (partes2)
    
    


    
    if secuencia:

        r1 = int(secuencia[:5])
        f1 = secuencia[5:8]
    
        r2 = int(secuencia[8:13])
        f2 = secuencia[13:16]
        
        if f1 == 'DIC':    
            funcionario =Partes.objects.filter( user_id=r1)
        else:
            funcionario =Partes.objects.filter( user_id=r2)
        
        print(funcionario)
        a2 = funcionario.first()
        print(a2)
        a3 = (a2.id )
        print (a3)
        fun=Partes.objects.get(pk=a3)
        r3 = int(secuencia[16:21])
        f3 = secuencia[21:24]
        
        r4 = int(secuencia[24:29])
        f4 = secuencia[29:32]
    
        r5 = int(secuencia[32:37])
        f5 = secuencia[37:40]
    
        r6 = int(secuencia[40:45])
        f6 = secuencia[45:48]

        print(r1)
        print(f1)
        print(r2)
        print(f2)
        print(r3)
        print(f3)
        print(r4)
        print(f4)
        print(r5)
        print(f5)
        print(r6)
        print(f6)
    else:
        print ("secuencia no asignada")    

    
    if request.method == "GET":
        enc = Contratos.objects.filter(pk=contrato_id).first()
        if not enc:
            encabezado = {
                'id':"",
                'uc_id':"",
                'tipocontrato':1,
                'datecontrato':datetime.today(),
                'datecontrato_ini':"",
                'datecontrato_fin':"",
                'parte1':546,
                'enCalidadDe1':'"CLIENTE"',
                'parte2':"",
                'enCalidadDe2':'"PRESTADOR DE SERVICIOS"',
                'lugarContrato':"",
                'ciudadContrato':"Mérida",
                'estadoContrato':"Yucatán",
                'paisContrato':"México",  
                'importeContrato':0.00,
                'npContrato':"",
                'imppContrato':0.00,
                'vhppContrato':285.00,
                'totalhorasContrato':"",
                'testigoContrato1':"",
                'testigoContrato2':"",
                'versionContrato':"",
        
                'status':"CAP",
                'rcap':xUsuario,
                
               
                'rstep1':r1,
                'rstep2':r2,
                'rstep3':r3,
                'rstep4':r4,
                'rstep5':r5,
                'rstep6':r6,
                
                
                'astep1':f1,
                'astep2':f2,
                'astep3':f3,
                'astep4':f4,
                'astep5':f5,
                'astep6':f6,
            
                'devuelto_por':"",

            
            }
            detalle=None
        else:
            
            encabezado = {
                'id':enc.id,
                'uc_id':enc.uc_id,
                'tipocontrato':enc.tipocontrato,
                'datecontrato':enc.datecontrato,
                'datecontrato_ini':enc.datecontrato_ini,
                'datecontrato_fin':enc.datecontrato_fin,
                'parte1':enc.parte1,
                'enCalidadDe1':enc.enCalidadDe1,
                'parte2':enc.parte2,
                'enCalidadDe2':enc.enCalidadDe2,
                'lugarContrato':enc.lugarContrato,
                'ciudadContrato':enc.ciudadContrato,
                'estadoContrato':enc.estadoContrato,
                'paisContrato':enc.paisContrato,  
                'importeContrato':enc.importeContrato,
                'npContrato':enc.npContrato,
                'imppContrato':enc.imppContrato,
                'vhppContrato':enc.vhppContrato,
                'totalhorasContrato':enc.totalhorasContrato,
                'testigoContrato1':enc.testigoContrato1,
                'testigoContrato2':enc.testigoContrato2,
                'versionContrato':enc.versionContrato,
        
                'status':enc.status,
                'rcap':enc.rcap,
                
               
                'rstep1':enc.rstep1,
                'rstep2':enc.rstep2,
                'rstep3':enc.rstep3,
                'rstep4':enc.rstep4,
                'rstep5':enc.rstep5,
                'rstep6':enc.rstep6,
                
                
                'astep1':enc.astep1,
                'astep2':enc.astep2,
                'astep3':enc.astep3,
                'astep4':enc.astep4,
                'astep5':enc.astep5,
                'astep6':enc.astep6,
            
                'devuelto_por':enc.devuelto_por,
                
                

            }

     
        
        detalle=Doctos.objects.filter(contrato=enc)
       
        contexto = {"requi":requisitos,"fun":funcionario,"fun2":partes2,"fun3":partes3,"enc":encabezado,"det":detalle,"departamentos":departamentos,"funcionarios":partes,"departamentos2":departamentos2,}
        
        return render(request,template_name,contexto)
    
   
    
    if request.method == "POST":
        
        tipocontrato=1
        datecontrato  = request.POST.get("datecontrato")
        datecontrato_ini  = request.POST.get("enc_datecontrato_ini")
        datecontrato_fin  = request.POST.get("enc_datecontrato_fin")
        
        parte1 = 546
        parte2 = request.POST.get("enc_nombreParte")

        enCalidadDe1 = "'CLIENTE '"
        enCalidadDe2 = "'PRESTADOR DE SERVICIOS '"        
        
        lugarContrato = request.POST.get("lugarContrato")
        ciudadContrato = "Mérida"
        estadoContrato = "Yucatán"
        paisContrato = "México"
        importeContrato = request.POST.get("enc_importeContrato")
        npContrato = request.POST.get("enc_npContrato")
        imppContrato = request.POST.get("enc_imppContrato")
        vhppContrato = 285
        totalhorasContrato = request.POST.get("enc_totalhorasContrato")
        testigoContrato1 = request.POST.get("enc_testigoContrato1")
        testigoContrato2 = request.POST.get("enc_testigoContrato2")
        versionContrato = request.POST.get("versionContrato")
        status = request.POST.get("status")
       
        fun=Partes.objects.get(pk=a3)
        tip=Tipocontrato.objects.get(pk=1)
        
        if parte2:
         suj=Partes.objects.get(pk=parte2)
        

        if not contrato_id:
            enc = Contratos(
                tipocontrato = tip,
                datecontrato  = datecontrato,
                datecontrato_ini  = datecontrato_ini,
                datecontrato_fin  = datecontrato_fin,
                
                parte1 = parte1,
                parte2 = suj,
                ciudadContrato = ciudadContrato,
                estadoContrato = estadoContrato,
                paisContrato = paisContrato,
                              
                enCalidadDe1 = enCalidadDe1,
                enCalidadDe2 = enCalidadDe2,

                importeContrato = importeContrato,
                npContrato = npContrato,
                imppContrato = imppContrato,
                vhppContrato = vhppContrato,
                totalhorasContrato = totalhorasContrato,
                testigoContrato1 = testigoContrato1,
                testigoContrato2 = testigoContrato2,

                status = status,
                rcap = xUsuario,
                current_user = xUsuario,
                rstep1 = r1,
                rstep2 = r2,
                rstep3 = r3,
                rstep4 = r4,
                rstep5 = r5,
                rstep6 = r6,
                
                astep1 = f1,
                astep2 = f2,
                astep3 = f3,
                astep4 = f4,
                astep5 = f5,
                astep6 = f6,
                devuelto_por = False,
            )
            if enc:
                enc.save()
                contrato_id = enc.id
        else:
            enc = Contratos.objects.filter(pk=contrato_id).first()
            if enc:
                
                enc.funcionario = fun
                
                enc.save()
                
        if not id:
            messages.error(request,'No Puedo Continuar No Pude Detectar No. de Contrato')
            return redirect("cto:contrato_list")
        
        documento = request.POST.get("documento")
        comentarioDocto = request.POST.get("comentarioDocto")
        req = Requisitos.objects.get(pk=documento)
        vigenciaFinDocto = request.POST.get("enc_vigenciaFinDocto")
        
        pdf= request.POST.get('pdf2')
        
        uploaded_file = request.FILES['pdf']
        print(uploaded_file)
            
        fs = FileSystemStorage()
        print(fs)
        name = fs.save(uploaded_file.name, uploaded_file)
        print(name)
        pdf = name

        if vigenciaFinDocto != "":
        
            det = Doctos(
                contrato=enc,
                documento=req,
                comentarioDocto=comentarioDocto,
                pdf=pdf,
                vigenciaFinDocto=vigenciaFinDocto,
            
            )
        
        if vigenciaFinDocto == "":
            
            det = Doctos(
                contrato=enc,
                documento=req,
                comentarioDocto=comentarioDocto,
                pdf=pdf,
              
            
            )
        
        if det:
            det.save()
            
        
            return redirect("cto:contrato_edit",contrato_id=contrato_id)

    return render(request,template_name,contexto)

class ContratosEdit(VistaBaseEdit):
        model = Contratos
        template_name="cto/contrato_form.html"
        form_class = ContratosForm
        success_url= reverse_lazy('cto:contrato_list')
        permission_required="cto.change_contratos"
        context_object_name = 'obj'





login_required(login_url="/login/")
@permission_required("cto.change_contratos",login_url="/login/")
def coverletter_export(request,id):
    
    contratos = Contratos.objects.filter(pk=id).first()  # Contrato en curso
    tipoc = Tipocontrato.objects.get(id=1) # Información del tipo de contrato
    valic = Valida.objects.filter(tipocontrato_id=1) # Validacion de información completa
    partes = Partes.objects.get(id=contratos.parte2_id) # Datos del contratado **sujeto del contrato
    patron = Partes.objects.get(id=546)  # Datos del contratante
    secue = Secuencia.objects.get(id=1) # Primer parrafo del contrato
    regimen = Regimen.objects.get(id=partes.regfiscalParte_id) # Régimen fiscal del contratado
    replegal = Partes.objects.get(id=549) # Datos del contratante
    letras = numero_to_letras(contratos.importeContrato.amount)
    pagolet = numero_to_letras(contratos.imppContrato.amount)
    currency = "${:,.2f}".format(contratos.importeContrato.amount)
    currency2 = "${:,.2f}".format(contratos.imppContrato.amount)
    document = Document()
    locale.setlocale(locale.LC_TIME, "es-MX")
    
    
   
        #set up font
    font = document.styles['Normal'].font
    font.name = 'Calibri'
    font.bold = True
    font.size = Pt(16)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri'

   

    
    # set up margins
    sections = document.sections
    for section in sections:
        
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.75)
        section.left_margin = Cm(2.00)
        section.right_margin = Cm(2.00)
        section.header_distance = Cm(0)
    
    header = document.sections[0].header
    htable=header.add_table(1, 2, Inches(5))
    
    for row in htable.rows:
        row.height = Inches(1.0)
    
    
    
    
    htable.alignment = WD_TABLE_ALIGNMENT.LEFT
    htab_cells=htable.rows[0].cells
    
    
    ht0=htab_cells[0].add_paragraph()
    paragraph_format = ht0.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    
    kh=ht0.add_run()
    
    THIS_FOLDER = os.path.dirname(os.path.abspath(__file__))
    my_file = os.path.join(THIS_FOLDER, 'logo.png')
    
    kh.add_picture(my_file, width=Inches(1.00))
    ht0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    ht1=htab_cells[1].add_paragraph("ESCUELA MODELO, S.C.P.")
    htable.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    htable.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER   
    
  
    # Get the user's fullname
    #if request.user.get_full_name():
    #   document_data_full_name = request.user.get_full_name()
    #else:
    #    document_data_full_name = "[NOMBRE] [APELLIDOS]"

    
    p001 = document.add_paragraph()
    p001.add_run(tipoc.tituloContrato, style = 'CommentsStyle')
    p001.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    p002 = document.add_paragraph()
    
    
    textox = tipoc.textoinicialContrato
    xnombreParte = partes.tituloParte + " " + partes.nombreParte
    xcurp = partes.curp[10:11]
    if xcurp == "H":
        xelolaParte = "EL"
    else:    
        xelolaParte = "LA"

    
    xenCalidadDe2 = contratos.enCalidadDe2
    xenCalidadDe2 = textox.replace("@enCalidadDe2" , xenCalidadDe2 )
    textox = xenCalidadDe2
    xenCalidadDe1 = textox.replace("@enCalidadDe1" , contratos.enCalidadDe1 )
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@elolaParte" , xelolaParte )
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@nombreParte2" , xnombreParte )
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@nombreParte" , patron.nombreParte )
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@tituloParteRL" , replegal.tituloParte )
    textox = xenCalidadDe1
    xenCalidadDe1 = textox.replace("@idrep_legalParte" , replegal.nombreParte )
    
    p002.add_run(xenCalidadDe1, style = 'CommentsStyle').bold = False
    p002.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p003 = document.add_paragraph()
    p003.add_run("CLÁUSULAS", style = 'CommentsStyle').bold = True
    p003.alignment = WD_ALIGN_PARAGRAPH.CENTER

   
    secue = Secuencia.objects.get(id=1)
    p004 = document.add_paragraph()
    textosecuex = secue.identificador + ".- " + secue.textoSecuencia
    textosecuex = textosecuex.replace("@enCalidadDe2" , contratos.enCalidadDe2 )
    p004.add_run(textosecuex, style = 'CommentsStyle').bold = True
    paragraph_format = p004.paragraph_format
    paragraph_format.left_indent = Inches(0.0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    
    
   
    
    
    nums = { 2, 3, 4, 5 , 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36 }

    
    secue = Secuencia.objects.get(id=2)
   
    for n in nums:
        secue = Secuencia.objects.get(id=n)   
        if secue.nivel2 == 0:
            
               secue = Secuencia.objects.get(id=n)
               p005 = document.add_paragraph()
               textosecue = secue.identificador + ".- " + secue.textoSecuencia
               textosecue = textosecue.replace("@enCalidadDe1" , contratos.enCalidadDe1 )
               textosecue = textosecue.replace("@enCalidadDe2" , contratos.enCalidadDe2 )
               textosecue = textosecue.replace("@datecontrato_ini" , contratos.datecontrato_ini.strftime("%d de %B de %Y"))
               textosecue = textosecue.replace("@datecontrato_fin" , contratos.datecontrato_fin.strftime("%d de %B de %Y"))
               textosecue = textosecue.replace("@FechaContrato" , contratos.datecontrato.strftime("%d de %B de %Y"))
               p005.add_run(textosecue, style = 'CommentsStyle').bold = True
               paragraph_format = p005.paragraph_format
               paragraph_format.space_before = Pt(0)
               paragraph_format.space_after = Pt(0)
               paragraph_format.left_indent = Inches(0.0)
               if len(textosecue) > 100:
                  paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
               else:   
                  paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
               
        else:
           
            secue = Secuencia.objects.get(id=n)
            p006= document.add_paragraph()
            textosecue = secue.identificador + ".- " + secue.textoSecuencia
            textosecue = textosecue.replace("@curp" , partes.curp )
            textosecue = textosecue.replace("@titulo_profParte" , partes.titulo_profParte )
            textosecue = textosecue.replace("@universidadParte" , partes.universidadParte )
            textosecue = textosecue.replace("@cedula_profParte" , partes.cedula_profParte )
            textosecue = textosecue.replace("@rfc" , partes.rfc )
            textosecue = textosecue.replace("@regfiscalParte" , regimen.nombreRegimen )
            textosecue = textosecue.replace("@domicilioParte" , partes.domicilioParte )
            textosecue = textosecue.replace("@enCalidadDe1" , contratos.enCalidadDe1 )
            textosecue = textosecue.replace("@enCalidadDe2" , contratos.enCalidadDe2 )
            textosecue = textosecue.replace("@domicilioPatron" , patron.domicilioParte )
            textosecue = textosecue.replace("@importeContrato" , currency )
            textosecue = textosecue.replace("@letras" , letras )
            textosecue = textosecue.replace("@npContrato" , str(contratos.npContrato) )
            textosecue = textosecue.replace("@imppContrato" , currency2 )
            textosecue = textosecue.replace("@pagolet" , pagolet )
            p006.add_run(textosecue, style = 'CommentsStyle').bold = False
            paragraph_format = p006.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.left_indent = Inches(0.4)
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    dtable=document.add_table(rows=2, cols=2)
    #dtable.style = "TableGrid"
    dtable.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    dtable.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER   
    
    
    dtab_cells=dtable.rows[0].cells
    dt1=dtab_cells[0].add_paragraph()
    dt1.add_run(contratos.enCalidadDe1, style = 'CommentsStyle').bold = True
    dt1.add_run( "_________________________")
    dt2=dtab_cells[1].add_paragraph()
    dt2.add_run(contratos.enCalidadDe2, style = 'CommentsStyle').bold = True
    dt2.add_run( "_________________________")
    paragraph_format = dt1.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = dt2.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    
    
    
    # Print the user's name
    #document_elements_heading = document.add_heading(document_data_full_name, 0)
    #document_elements_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
 
    # Print biography and careerpath
    
    # Add empty paragraph
    #document.add_paragraph()
 
    # Sincerely and name
    #document.add_paragraph("Atentamente,\n" + document_data_full_name) 

    
    
    document_data = io.BytesIO()
    print(document_data)
    document.save(document_data)
    document_data.seek(0)
    response = HttpResponse(
        document_data.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = 'attachment; filename = "Contrato.docx"'
    response["Content-Encoding"] = "UTF-8"
    return response
                                                                                                                                                        

@login_required(login_url="/login/")
@permission_required("cto.change_contratos",login_url="/login/")
def contratosAvanza(request,id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.method=="POST":
        
        if not contratos.fcap:
           contratos.fcap = datetime.today()
           contratos.current_user = contratos.rstep1
           contratos.status = contratos.astep1
           contratos.devuelto_por = contratos.rcap
           contratos.save()
           return HttpResponse("OK")
        #return HttpResponse("FAIL")
        if not contratos.fstep1:
           contratos.fstep1 = datetime.today()
           contratos.current_user = contratos.rstep2
           contratos.status = contratos.astep2
           contratos.save()
           return HttpResponse("OK")
        if not contratos.fstep2:
           contratos.fstep2 = datetime.today()
           contratos.current_user = contratos.rstep3
           contratos.status = contratos.astep3
           contratos.save()
           return HttpResponse("OK")
        if not contratos.fstep3:
           contratos.fstep3 = datetime.today()
           contratos.current_user = contratos.rstep4
           contratos.status = contratos.astep4
           contratos.save()
           return HttpResponse("OK")
        if not contratos.fstep4:
           contratos.fstep4 = datetime.today()
           contratos.current_user = contratos.rstep5
           contratos.status = contratos.astep5
           contratos.save()
           return HttpResponse("OK")
        if not contratos.fstep5:
           contratos.fstep5 = datetime.today()
           contratos.current_user = contratos.rstep6
           contratos.status = contratos.astep6
           contratos.save()
           return HttpResponse("OK")
        
        if not contratos.fstep6:
           contratos.fstep6 = datetime.today()
           #contratos.current_user = contratos.rstep6
           contratos.status = "FIN"
           contratos.save()
           return HttpResponse("OK")
        return HttpResponse("FAIL")
    
    
    return HttpResponse("FAIL") 

@login_required(login_url="/login/")
@permission_required("cto.change_contratos",login_url="/login/")
def contratosDevuelve(request,id):
    contratos = Contratos.objects.filter(pk=id).first()

    if request.method=="POST":
        if not contratos.fstep6:
           contratos.fcap = contratos.fstep6
           contratos.fstep5 = contratos.fstep6
           contratos.fstep4 = contratos.fstep6
           contratos.fstep3 = contratos.fstep6
           contratos.fstep2 = contratos.fstep6
           contratos.fstep1 = contratos.fstep6
           contratos.devuelto_por = contratos.current_user
           contratos.current_user = contratos.rcap
           contratos.status = "CAP"
           
           contratos.save()
           return HttpResponse("OK")
        return HttpResponse("FAIL")
    return HttpResponse("FAIL")

def numero_to_letras(numero):
    
    indicador = [
        ("",""),("MIL","MIL"),("MILLON","MILLONES"),("MIL","MIL"),("BILLON","BILLONES")]
	
    entero = int(numero)
	
    decimal = int(round((numero - entero)*100))

    contador = 0
	
    numero_letras = ""

    while entero >0:
        
        a = entero % 1000
		
        if contador == 0:
		
        	en_letras = convierte_cifra(a,1).strip()
		
        else :
		
        	en_letras = convierte_cifra(a,0).strip()
		
        if a==0:
		
        	numero_letras = en_letras+" "+numero_letras
		

        
        elif a==1:
		
        	if contador in (1,3):
		
        		numero_letras = indicador[contador][0]+" "+numero_letras
		
        	else:
		
        		numero_letras = en_letras+" "+indicador[contador][0]+" "+numero_letras
		
        else:
		
        	numero_letras = en_letras+" "+indicador[contador][1]+" "+numero_letras
		
        numero_letras = numero_letras.strip()
		
        contador = contador + 1
		
        entero = int(entero / 1000)
    
    numero_letras = numero_letras+" PESOS " + str(decimal) +"/100 M.N."
    return (numero_letras)

def convierte_cifra(numero,sw):
    
    lista_centana = ["",("CIEN","CIENTO"),"DOSCIENTOS","TRESCIENTOS","CUATROCIENTOS","QUINIENTOS","SEISCIENTOS","SETECIENTOS","OCHOCIENTOS","NOVECIENTOS"]
	
    lista_decena = ["",("DIEZ","ONCE","DOCE","TRECE","CATORCE","QUINCE","DIECISEIS","DIECISIETE","DIECIOCHO","DIECINUEVE"),
					("VEINTE","VEINTI"),("TREINTA","TREINTA Y"),("CUARENTA" , "CUARENTA Y"),
					("CINCUENTA" , "CINCUENTA Y"),("SESENTA" , "SESENTA Y"),
					("SETENTA" , "SETENTA Y"),("OCHENTA" , "OCHENTA Y"),
					("NOVENTA" , "NOVENTA Y")
				]
	
    lista_unidad = ["",("UN","UNO"),"DOS","TRES","CUATRO","CINCO","SEIS","SIETE","OCHO","NUEVE"]
	
    centena = int (numero / 100)
	
    decena = int((numero -(centena * 100))/10)

    unidad = int(numero - (centena * 100 + decena * 10))

    texto_centena = ""
	
    texto_decena = ""
	
    texto_unidad = ""

    texto_centena = lista_centana[centena]
	
    if centena == 1:
	
    	if (decena + unidad)!=0:
	
    		texto_centena = texto_centena[1]
	
    	else :
	
     		texto_centena = texto_centena[0]
    
    texto_decena = lista_decena[decena]
	
    if decena == 1 :
	
    	 texto_decena = texto_decena[unidad]
 	
    elif decena > 1 :
 	
     	if unidad != 0 :
 	
     		texto_decena = texto_decena[1]
 	
     	else:
 	
     		texto_decena = texto_decena[0] 

    if decena != 1:
     		
            texto_unidad = lista_unidad[unidad]
 		
            if unidad == 1:
 		
         	    texto_unidad = texto_unidad[sw]

    return "%s %s %s" %(texto_centena,texto_decena,texto_unidad)